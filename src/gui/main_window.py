"""
Main application window
"""

from PyQt6.QtWidgets import (QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
                              QPushButton, QLabel, QFileDialog, QMessageBox,
                              QTextEdit, QTableWidget, QTableWidgetItem, QSplitter,
                              QLineEdit, QComboBox, QGroupBox, QCheckBox, QTabWidget,
                              QStyledItemDelegate, QDialog, QTextBrowser)
from PyQt6.QtCore import Qt, QSortFilterProxyModel, QTimer, QUrl, QEvent
from PyQt6.QtGui import QStandardItemModel, QStandardItem, QDesktopServices, QPixmap, QCursor
import sys
import os
import logging

logger = logging.getLogger(__name__)


def get_resource_path(relative_path):
    """Get absolute path to resource, works for dev and for PyInstaller"""
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        # If not running as PyInstaller bundle, use the project root
        base_path = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

    return os.path.join(base_path, relative_path)

# Add parent directory to path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from parser.xml_parser import ProfileParser
from parser.label_generator import LabelGenerator
from models.profile_model import ControlProfile
from utils.settings import AppSettings
from utils.version import get_version
from utils.device_splitter import get_device_for_input

# Try to import PDF widget - prefer QtPdf (native, lightweight) over WebEngine
PDF_WIDGET_AVAILABLE = False
PDF_WIDGET_CLASS = None

try:
    from gui.qtpdf_device_widget import QtPdfDeviceWidget
    PDF_WIDGET_CLASS = QtPdfDeviceWidget
    PDF_WIDGET_AVAILABLE = True
    logger.info("Using QtPdf for device view (native Qt PDF viewer)")
except ImportError as e:
    logger.warning(f"QtPdf widget not available: {e}")
    # Fallback to WebEngine if QtPdf not available
    try:
        from gui.webengine_pdf_widget import WebEnginePDFWidget
        PDF_WIDGET_CLASS = WebEnginePDFWidget
        PDF_WIDGET_AVAILABLE = True
        logger.info("Using WebEngine for device view (Chromium-based)")
    except ImportError as e2:
        PDF_WIDGET_AVAILABLE = False
        logger.warning(f"WebEngine also not available: {e2}")
        logger.error("No PDF widget available - Device View will be disabled")


class SelectAllDelegate(QStyledItemDelegate):
    """Custom delegate - clean editing with proper text replacement"""

    def __init__(self, parent, main_window):
        super().__init__(parent)
        self.main_window = main_window

    def createEditor(self, parent, option, index):
        """Create editor with proper configuration"""
        editor = QLineEdit(parent)
        # Force solid white background and black text
        editor.setStyleSheet("QLineEdit { background-color: white; color: black; }")
        return editor

    def setEditorData(self, editor, index):
        """Set editor data and select all for easy replacement"""
        if isinstance(editor, QLineEdit):
            # Get current text
            text = index.model().data(index, Qt.ItemDataRole.DisplayRole)
            editor.setText(text or "")
            # Ensure editor has focus before selecting
            editor.setFocus()
            # Use a small delay to ensure editor is ready
            QTimer.singleShot(0, editor.selectAll)
        else:
            super().setEditorData(editor, index)


class MainWindow(QMainWindow):
    """Main application window for SC Profile Editor"""

    def __init__(self):
        super().__init__()
        self.version = get_version()
        self.setWindowTitle(f"Star Citizen Profile Editor v{self.version}")
        self.setGeometry(100, 100, 1200, 800)

        # Initialize settings manager
        self.settings = AppSettings()

        # Store current profile
        self.current_profile = None
        self.current_profile_path = None

        # Store all bindings for filtering
        self.all_bindings = []

        # Create UI
        self.setup_ui()

        # Restore window geometry if saved
        self.restore_window_state()

        # Auto-load last profile after UI is shown (use QTimer to defer until after show())
        QTimer.singleShot(100, self.auto_load_last_profile)

    def setup_ui(self):
        """Set up the user interface"""
        # Central widget
        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        # Main layout
        main_layout = QVBoxLayout()
        central_widget.setLayout(main_layout)

        # Header
        header_layout = QHBoxLayout()
        title_label = QLabel(f"Star Citizen Profile Editor v{self.version}")
        title_label.setStyleSheet("font-size: 24px; font-weight: bold; margin: 10px;")
        header_layout.addWidget(title_label)
        header_layout.addStretch()

        # Import button (leftmost)
        import_btn = QPushButton("Import Profile XML")
        import_btn.setStyleSheet("padding: 10px 20px; font-size: 14px; background-color: #4CAF50; color: white;")
        import_btn.clicked.connect(self.import_profile)
        header_layout.addWidget(import_btn)

        # Export buttons
        self.export_csv_btn = QPushButton("Export CSV")
        self.export_csv_btn.setStyleSheet("padding: 10px 20px; font-size: 14px;")
        self.export_csv_btn.clicked.connect(self.export_csv)
        self.export_csv_btn.setEnabled(False)
        header_layout.addWidget(self.export_csv_btn)

        self.export_pdf_btn = QPushButton("Export PDF")
        self.export_pdf_btn.setStyleSheet("padding: 10px 20px; font-size: 14px;")
        self.export_pdf_btn.clicked.connect(self.export_pdf)
        self.export_pdf_btn.setEnabled(False)
        header_layout.addWidget(self.export_pdf_btn)

        self.export_word_btn = QPushButton("Export Word")
        self.export_word_btn.setStyleSheet("padding: 10px 20px; font-size: 14px;")
        self.export_word_btn.clicked.connect(self.export_word)
        self.export_word_btn.setEnabled(False)
        header_layout.addWidget(self.export_word_btn)

        self.export_graphic_btn = QPushButton("Export Graphic")
        self.export_graphic_btn.setStyleSheet("padding: 10px 20px; font-size: 14px;")
        self.export_graphic_btn.clicked.connect(self.export_graphic)
        self.export_graphic_btn.setEnabled(False)
        header_layout.addWidget(self.export_graphic_btn)

        # Save Profile button (for remapped controls)
        self.save_profile_btn = QPushButton("💾 Save Profile")
        self.save_profile_btn.setStyleSheet("padding: 10px 20px; font-size: 14px; background-color: #2196F3; color: white;")
        self.save_profile_btn.clicked.connect(self.save_profile)
        self.save_profile_btn.setEnabled(False)  # Disabled until profile is modified
        header_layout.addWidget(self.save_profile_btn)

        # Help button (rightmost)
        help_btn = QPushButton("Help")
        help_btn.setStyleSheet("padding: 10px 20px; font-size: 14px;")
        help_btn.clicked.connect(self.show_help)
        header_layout.addWidget(help_btn)

        main_layout.addLayout(header_layout)

        # Profile info label
        self.profile_info_label = QLabel("No profile loaded")
        self.profile_info_label.setStyleSheet("font-size: 12px; margin: 5px; color: #666;")
        main_layout.addWidget(self.profile_info_label)

        # Filter toolbar
        filter_group = QGroupBox("Filters")
        filter_layout = QHBoxLayout()
        filter_group.setLayout(filter_layout)

        # Search box
        filter_layout.addWidget(QLabel("Search:"))
        self.search_box = QLineEdit()
        self.search_box.setPlaceholderText("Type to filter actions, inputs, or devices...")
        self.search_box.textChanged.connect(self.apply_filters)
        filter_layout.addWidget(self.search_box, 2)

        # Device filter
        filter_layout.addWidget(QLabel("Device:"))
        self.device_filter = QComboBox()
        self.device_filter.addItem("All Devices")
        self.device_filter.currentTextChanged.connect(self.apply_filters)
        filter_layout.addWidget(self.device_filter, 1)

        # Action Map filter
        filter_layout.addWidget(QLabel("Action Map:"))
        self.actionmap_filter = QComboBox()
        self.actionmap_filter.addItem("All Action Maps")
        self.actionmap_filter.currentTextChanged.connect(self.apply_filters)
        filter_layout.addWidget(self.actionmap_filter, 1)

        # Hide unmapped keys checkbox
        self.hide_unmapped_checkbox = QCheckBox("Hide Unmapped Keys")
        self.hide_unmapped_checkbox.stateChanged.connect(self.apply_filters)
        filter_layout.addWidget(self.hide_unmapped_checkbox)

        # Show detailed checkbox
        self.show_detailed_checkbox = QCheckBox("Show Detailed")
        self.show_detailed_checkbox.stateChanged.connect(self.toggle_detailed_view)
        filter_layout.addWidget(self.show_detailed_checkbox)

        # Clear filters button
        clear_btn = QPushButton("Clear Filters")
        clear_btn.clicked.connect(self.clear_filters)
        filter_layout.addWidget(clear_btn)

        main_layout.addWidget(filter_group)

        # Create tab widget for different views
        self.tab_widget = QTabWidget()
        main_layout.addWidget(self.tab_widget)

        # Tab 1: Controls Table View
        table_tab = QWidget()
        table_layout = QVBoxLayout()
        table_tab.setLayout(table_layout)

        # Splitter for info and table
        splitter = QSplitter(Qt.Orientation.Vertical)

        # Profile summary text
        self.profile_summary = QTextEdit()
        self.profile_summary.setReadOnly(True)
        self.profile_summary.setMaximumHeight(150)
        self.profile_summary.setPlaceholderText("Profile summary will appear here...")
        splitter.addWidget(self.profile_summary)

        # Controls table
        self.controls_table = QTableWidget()
        self.controls_table.setColumnCount(6)  # Max columns for detailed view
        self.controls_table.setHorizontalHeaderLabels([
            "Action Map", "Action", "Action (Override)", "Input Code", "Input Label", "Device"
        ])
        self.controls_table.horizontalHeader().setStretchLastSection(False)
        self.controls_table.setAlternatingRowColors(True)
        self.controls_table.setSortingEnabled(True)
        self.controls_table.setEditTriggers(QTableWidget.EditTrigger.DoubleClicked)
        self.controls_table.itemChanged.connect(self.on_cell_edited)
        self.controls_table.itemDoubleClicked.connect(self.on_item_double_clicked)

        # Enable context menu for remapping
        self.controls_table.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.controls_table.customContextMenuRequested.connect(self.show_table_context_menu)

        # Set custom delegate for column 2 to handle text selection
        delegate = SelectAllDelegate(self.controls_table, self)
        self.controls_table.setItemDelegateForColumn(2, delegate)

        splitter.addWidget(self.controls_table)

        # Track editing state
        self._editing_item = None
        self._editing_default_text = None  # Stores the default label (without custom override)

        # Track editable columns (column 2 is Action Override, editable only in detailed view)
        self.editable_columns = {2}  # Column 2: Action (Override)

        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 3)

        table_layout.addWidget(splitter)
        self.tab_widget.addTab(table_tab, "Controls Table")

        # Tab 2: Device View (PDF-based device visualization) - Optional
        templates_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "visual-templates")
        if PDF_WIDGET_AVAILABLE:
            self.pdf_device_widget = PDF_WIDGET_CLASS(templates_dir)
            self.pdf_device_widget.export_available_changed.connect(self.export_graphic_btn.setEnabled)
            self.tab_widget.addTab(self.pdf_device_widget, "Device View")
        else:
            self.pdf_device_widget = None
            logger.warning("Device View tab disabled - no PDF widget available")

        # Tab 3: About
        about_tab = QWidget()
        about_layout = QVBoxLayout()
        about_tab.setLayout(about_layout)

        # About content in a scrollable text browser
        about_browser = QTextBrowser()
        about_browser.setOpenExternalLinks(True)

        about_html = f"""
        <html>
        <head>
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    padding: 20px;
                    line-height: 1.6;
                }}
                h1 {{
                    color: #2c3e50;
                    border-bottom: 3px solid #3498db;
                    padding-bottom: 10px;
                }}
                h2 {{
                    color: #34495e;
                    border-bottom: 2px solid #95a5a6;
                    padding-bottom: 5px;
                    margin-top: 30px;
                }}
                h3 {{
                    color: #555;
                    margin-top: 20px;
                }}
                p {{
                    margin: 10px 0;
                }}
                ul {{
                    margin-left: 20px;
                }}
                li {{
                    margin: 5px 0;
                }}
                .placeholder {{
                    color: #888;
                    font-style: italic;
                }}
            </style>
        </head>
        <body>
            <h1>Star Citizen Profile Editor v{self.version}</h1>

            <h2>About This Project</h2>
            <p class="placeholder">
                [This section will contain information about the project and Osiris DevWorks]
            </p>

            <h2>Acknowledgements</h2>
            <p>
                Special thanks to the following individuals for their valuable assistance
                and contributions to this project:
            </p>
            <ul>
                <li><strong>GurningBoose</strong></li>
                <li><strong>Hawkwar</strong></li>
                <li><strong>Nazgul-Five 'Maverick'</strong></li>
                <li><strong>Tichro 'BreakPoint'</strong></li>
                <li><strong>UntoldForce</strong></li>
            </ul>
        </body>
        </html>
        """

        about_browser.setHtml(about_html)
        about_layout.addWidget(about_browser)

        self.tab_widget.addTab(about_tab, "About")

        # Connect tab change signal to sync device selection
        self.tab_widget.currentChanged.connect(self.on_tab_changed)

        # Footer with Osiris DevWorks (left) and PayPal (right)
        footer_layout = QHBoxLayout()

        # Osiris DevWorks button (left side)
        self.osiris_button = QLabel()
        osiris_image_path = get_resource_path(os.path.join("assets", "osiris-devworks.png"))

        # Try to load Osiris image, fall back to text if not found
        if os.path.exists(osiris_image_path):
            pixmap = QPixmap(osiris_image_path)
            # Scale to reasonable size (max height 40px to accommodate the logo design)
            if pixmap.height() > 40:
                pixmap = pixmap.scaledToHeight(40, Qt.TransformationMode.SmoothTransformation)
            self.osiris_button.setPixmap(pixmap)
        else:
            # Fallback to styled text button
            self.osiris_button.setText("Osiris DevWorks")
            self.osiris_button.setStyleSheet("""
                QLabel {
                    background-color: #1a1f2e;
                    color: #c9a961;
                    padding: 8px 16px;
                    border-radius: 4px;
                    font-size: 12px;
                    font-weight: bold;
                }
                QLabel:hover {
                    background-color: #242938;
                }
            """)

        self.osiris_button.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self.osiris_button.mousePressEvent = self.open_discord_link
        footer_layout.addWidget(self.osiris_button)

        # Stretch to push donation buttons to the right
        footer_layout.addStretch()

        # Donation label
        donation_label = QLabel("Support this project:")
        donation_label.setStyleSheet("font-size: 12px; color: #666; margin-right: 5px;")
        footer_layout.addWidget(donation_label)

        # PayPal button (right side)
        self.paypal_button = QLabel()
        paypal_image_path = get_resource_path(os.path.join("assets", "paypal.png"))

        # Try to load PayPal image, fall back to text if not found
        if os.path.exists(paypal_image_path):
            paypal_pixmap = QPixmap(paypal_image_path)
            # Scale to match Osiris button (max height 40px)
            if paypal_pixmap.height() > 40:
                paypal_pixmap = paypal_pixmap.scaledToHeight(40, Qt.TransformationMode.SmoothTransformation)
            self.paypal_button.setPixmap(paypal_pixmap)
        else:
            # Fallback to styled text button
            self.paypal_button.setText("Donate via PayPal")
            self.paypal_button.setStyleSheet("""
                QLabel {
                    background-color: #0070ba;
                    color: white;
                    padding: 8px 16px;
                    border-radius: 4px;
                    font-size: 12px;
                    font-weight: bold;
                }
                QLabel:hover {
                    background-color: #005ea6;
                }
            """)

        self.paypal_button.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self.paypal_button.mousePressEvent = self.open_paypal_donation
        footer_layout.addWidget(self.paypal_button)

        # Spacer between PayPal and Venmo
        footer_layout.addSpacing(10)

        # Venmo button (right side)
        self.venmo_button = QLabel()
        venmo_image_path = get_resource_path(os.path.join("assets", "venmo.png"))

        # Try to load Venmo image, fall back to text button
        if os.path.exists(venmo_image_path):
            venmo_pixmap = QPixmap(venmo_image_path)
            # Scale to match Osiris button (max height 40px)
            if venmo_pixmap.height() > 40:
                venmo_pixmap = venmo_pixmap.scaledToHeight(40, Qt.TransformationMode.SmoothTransformation)
            self.venmo_button.setPixmap(venmo_pixmap)
        else:
            # Fallback to styled text button
            self.venmo_button.setText("Venmo")
            self.venmo_button.setStyleSheet("""
                QLabel {
                    background-color: #008CFF;
                    color: white;
                    padding: 8px 16px;
                    border-radius: 4px;
                    font-size: 12px;
                    font-weight: bold;
                }
                QLabel:hover {
                    background-color: #0074D9;
                }
            """)

        self.venmo_button.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self.venmo_button.mousePressEvent = self.open_venmo_donation
        footer_layout.addWidget(self.venmo_button)

        main_layout.addLayout(footer_layout)

        # Status bar
        self.statusBar().showMessage("Ready")

    def restore_window_state(self):
        """Restore saved window geometry and state"""
        geometry = self.settings.get_window_geometry()
        if geometry:
            self.restoreGeometry(geometry)
            logger.debug("Restored window geometry")

        state = self.settings.get_window_state()
        if state:
            self.restoreState(state)
            logger.debug("Restored window state")

    def closeEvent(self, event):
        """Save window state before closing"""
        self.settings.set_window_geometry(self.saveGeometry())
        self.settings.set_window_state(self.saveState())
        logger.debug("Saved window state")
        event.accept()

    def auto_load_last_profile(self):
        """Automatically load the last opened profile if it exists"""
        last_profile_path = self.settings.get_last_profile_path()

        if not last_profile_path:
            logger.info("No last profile found")
            return

        if not os.path.exists(last_profile_path):
            logger.warning(f"Last profile no longer exists: {last_profile_path}")
            self.settings.clear_last_profile_path()
            return

        logger.info(f"Auto-loading last profile: {last_profile_path}")
        self.load_profile_from_path(last_profile_path)

    def import_profile(self):
        """Handle profile import button click"""
        # Get last directory if available
        last_path = self.settings.get_last_profile_path()
        start_dir = os.path.dirname(last_path) if last_path and os.path.exists(last_path) else ""

        # Open file dialog
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select Star Citizen Profile XML",
            start_dir,
            "XML Files (*.xml);;All Files (*)"
        )

        if not file_path:
            return  # User cancelled

        self.load_profile_from_path(file_path)

    def load_profile_from_path(self, file_path: str):
        """
        Load a profile from the given file path

        Args:
            file_path: Path to the profile XML file
        """
        try:
            # Parse the profile
            self.statusBar().showMessage(f"Loading profile: {file_path}")
            logger.info(f"Loading profile: {file_path}")
            parser = ProfileParser(file_path)
            self.current_profile = parser.parse()
            self.current_profile_path = file_path

            # Save as last opened profile
            self.settings.set_last_profile_path(file_path)

            # Update UI with profile data
            self.display_profile()

            # Enable export buttons
            self.export_csv_btn.setEnabled(True)
            self.export_pdf_btn.setEnabled(True)
            self.export_word_btn.setEnabled(True)

            self.statusBar().showMessage(f"Successfully loaded: {self.current_profile.profile_name}")
            logger.info(f"Successfully loaded profile: {self.current_profile.profile_name}")

        except FileNotFoundError:
            QMessageBox.critical(self, "Error", f"File not found: {file_path}")
            self.statusBar().showMessage("Error: File not found")
            logger.error(f"File not found: {file_path}")
        except ValueError as e:
            QMessageBox.critical(self, "Parse Error", f"Failed to parse XML file:\n{str(e)}")
            self.statusBar().showMessage("Error: Failed to parse XML")
            logger.error(f"Parse error: {e}", exc_info=True)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"An unexpected error occurred:\n{str(e)}")
            self.statusBar().showMessage("Error loading profile")
            logger.error(f"Error loading profile: {e}", exc_info=True)

    def display_profile(self):
        """Display the loaded profile data"""
        if not self.current_profile:
            return

        profile = self.current_profile

        # Update profile info label
        self.profile_info_label.setText(f"Profile: {profile.profile_name}")

        # Build summary text
        summary_parts = []
        summary_parts.append(f"Profile Name: {profile.profile_name}")
        summary_parts.append(f"Devices: {len(profile.devices)}")

        # List devices
        for device in profile.devices:
            device_name = device.product_name if device.product_name else device.device_type.capitalize()
            summary_parts.append(f"  - {device.device_type.capitalize()} {device.instance}: {device_name}")

        summary_parts.append(f"\nAction Maps: {len(profile.action_maps)}")
        summary_parts.append(f"Total Bindings: {len(profile.get_all_bindings())}")

        if profile.categories:
            summary_parts.append(f"\nCategories: {', '.join(profile.categories[:5])}")
            if len(profile.categories) > 5:
                summary_parts.append(f"  ... and {len(profile.categories) - 5} more")

        self.profile_summary.setText('\n'.join(summary_parts))

        # Populate controls table
        self.populate_controls_table()

        # Apply filters to show the rows (important when loading a new profile)
        self.apply_filters()

        # Load profile into Device View widget
        if self.pdf_device_widget:
            self.pdf_device_widget.load_profile(profile)

    def populate_controls_table(self):
        """Populate the controls table with bindings"""
        if not self.current_profile:
            return

        # Disable sorting while populating
        self.controls_table.setSortingEnabled(False)

        # Block signals to prevent itemChanged from triggering during population
        self.controls_table.blockSignals(True)

        # Get all bindings and store for filtering
        self.all_bindings = []
        for action_map in self.current_profile.action_maps:
            for binding in action_map.actions:
                self.all_bindings.append((action_map.name, binding))

        # Populate filter dropdowns
        self.populate_filter_dropdowns()

        # Set row count
        self.controls_table.setRowCount(len(self.all_bindings))

        # Determine if we're in detailed view
        is_detailed = self.show_detailed_checkbox.isChecked()

        # Configure column visibility based on view mode
        if is_detailed:
            # Detailed view: Show all 6 columns
            # 0: Action Map, 1: Action (original), 2: Action (Override), 3: Input Code, 4: Input Label, 5: Device
            self.controls_table.setColumnHidden(1, False)
            self.controls_table.setColumnHidden(2, False)
            self.controls_table.setColumnHidden(3, False)
            self.controls_table.setColumnHidden(4, False)
        else:
            # Default view: Show 4 columns (0: Action Map, 2: Action, 4: Input Label, 5: Device)
            # Hide columns 1, 3
            self.controls_table.setColumnHidden(1, True)
            self.controls_table.setColumnHidden(3, True)
            self.controls_table.setColumnHidden(4, False)  # Show Input Label
            self.controls_table.setColumnHidden(2, False)  # Show override column (as "Action")

        # Populate table
        for row, (action_map_name, binding) in enumerate(self.all_bindings):
            # Column 0: Action Map
            action_map_label = LabelGenerator.generate_actionmap_label(action_map_name)
            action_map_item = QTableWidgetItem(action_map_label)
            action_map_item.setFlags(action_map_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.controls_table.setItem(row, 0, action_map_item)

            # Column 1: Action (original auto-generated) - only visible in detailed view
            action_label_original = LabelGenerator.generate_action_label(binding.action_name)
            action_original_item = QTableWidgetItem(action_label_original)
            action_original_item.setFlags(action_original_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.controls_table.setItem(row, 1, action_original_item)

            # Column 2: Action (with override) - editable, visible in both views
            action_label_override = LabelGenerator.get_action_label(binding.action_name, binding)
            action_override_item = QTableWidgetItem(action_label_override)
            # Store binding reference in the item for later retrieval during editing
            action_override_item.setData(Qt.ItemDataRole.UserRole, (action_map_name, binding))
            self.controls_table.setItem(row, 2, action_override_item)

            # Column 3: Input Code (raw) - only visible in detailed view
            input_code_item = QTableWidgetItem(binding.input_code)
            input_code_item.setFlags(input_code_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.controls_table.setItem(row, 3, input_code_item)

            # Column 4: Input Label (human-readable) - only visible in detailed view
            input_label = LabelGenerator.generate_input_label(binding.input_code)
            input_label_item = QTableWidgetItem(input_label)
            input_label_item.setFlags(input_label_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.controls_table.setItem(row, 4, input_label_item)

            # Column 5: Device (parsed from input code)
            device = self.parse_device_from_input(binding.input_code)
            device_item = QTableWidgetItem(device)
            device_item.setFlags(device_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.controls_table.setItem(row, 5, device_item)

        # Re-enable signals
        self.controls_table.blockSignals(False)

        # Re-enable sorting
        self.controls_table.setSortingEnabled(True)

        # Resize columns to content
        self.controls_table.resizeColumnsToContents()

        # Set column widths based on view mode
        if is_detailed:
            self.controls_table.setColumnWidth(0, 200)  # Action Map
            self.controls_table.setColumnWidth(1, 200)  # Action (Original)
            self.controls_table.setColumnWidth(2, 200)  # Action (Override)
            self.controls_table.setColumnWidth(3, 150)  # Input Code
            self.controls_table.setColumnWidth(4, 200)  # Input Label
            self.controls_table.setColumnWidth(5, 150)  # Device
        else:
            self.controls_table.setColumnWidth(0, 200)  # Action Map
            self.controls_table.setColumnWidth(2, 250)  # Action (Override)
            self.controls_table.setColumnWidth(4, 200)  # Input Label
            self.controls_table.setColumnWidth(5, 150)  # Device

    def parse_device_from_input(self, input_code: str) -> str:
        """Parse device type from input code"""
        if input_code.startswith('kb'):
            return "Keyboard"
        elif input_code.startswith('js'):
            # Extract joystick instance
            import re
            match = re.match(r'js(\d+)_', input_code)
            if match:
                instance = match.group(1)
                # Try to find device name
                for device in self.current_profile.devices:
                    if device.device_type == 'joystick' and device.instance == int(instance):
                        device_name = device.product_name if device.product_name else f"Joystick {instance}"
                        # Split composite devices (e.g., VKB Gladiator + SEM)
                        return get_device_for_input(device_name, input_code)
                return f"Joystick {instance}"
        elif 'mouse' in input_code.lower():
            return "Mouse"
        return "Unknown"

    def populate_filter_dropdowns(self):
        """Populate the device and action map filter dropdowns"""
        # Get unique devices
        devices = set()
        action_maps = set()

        for action_map_name, binding in self.all_bindings:
            device = self.parse_device_from_input(binding.input_code)
            devices.add(device)
            action_map_label = LabelGenerator.generate_actionmap_label(action_map_name)
            action_maps.add(action_map_label)

        # Save current selections before clearing
        current_device = self.device_filter.currentText()
        current_actionmap = self.actionmap_filter.currentText()

        # Update device filter
        self.device_filter.blockSignals(True)
        self.device_filter.clear()
        self.device_filter.addItem("All Devices")
        for device in sorted(devices):
            self.device_filter.addItem(device)

        # Restore previous selection if it still exists
        if current_device:
            index = self.device_filter.findText(current_device)
            if index >= 0:
                self.device_filter.setCurrentIndex(index)

        self.device_filter.blockSignals(False)

        # Update action map filter
        self.actionmap_filter.blockSignals(True)
        self.actionmap_filter.clear()
        self.actionmap_filter.addItem("All Action Maps")
        for action_map in sorted(action_maps):
            self.actionmap_filter.addItem(action_map)

        # Restore previous selection if it still exists
        if current_actionmap:
            index = self.actionmap_filter.findText(current_actionmap)
            if index >= 0:
                self.actionmap_filter.setCurrentIndex(index)

        self.actionmap_filter.blockSignals(False)

    def apply_filters(self):
        """Apply current filter settings to the table"""
        if not self.all_bindings:
            return

        search_text = self.search_box.text().lower()
        device_filter = self.device_filter.currentText()
        actionmap_filter = self.actionmap_filter.currentText()
        hide_unmapped = self.hide_unmapped_checkbox.isChecked()

        # Show/hide rows based on filters
        for row in range(self.controls_table.rowCount()):
            show_row = True

            # Search filter - check all columns
            if search_text:
                row_text = ""
                for col in range(self.controls_table.columnCount()):
                    item = self.controls_table.item(row, col)
                    if item:
                        row_text += item.text().lower() + " "

                if search_text not in row_text:
                    show_row = False

            # Device filter
            if show_row and device_filter != "All Devices":
                device_item = self.controls_table.item(row, 5)  # Device is now column 5
                if device_item and device_item.text() != device_filter:
                    show_row = False

            # Action map filter
            if show_row and actionmap_filter != "All Action Maps":
                actionmap_item = self.controls_table.item(row, 0)
                if actionmap_item and actionmap_item.text() != actionmap_filter:
                    show_row = False

            # Unmapped keys filter - check if input_code == input_label (after stripping)
            # or if the label indicates an empty/unmapped binding
            if show_row and hide_unmapped:
                input_code_item = self.controls_table.item(row, 3)  # Input Code is now column 3
                input_label_item = self.controls_table.item(row, 4)  # Input Label is now column 4
                if input_code_item and input_label_item:
                    input_code = input_code_item.text().strip()
                    input_label = input_label_item.text().strip()

                    # Check if they're the same after stripping
                    if input_code == input_label:
                        show_row = False
                    # Also check for patterns that indicate unmapped keys
                    # e.g., "Keyboard: ", "Joystick 1: ", "Mouse: " with nothing after
                    elif input_label.endswith(': ') or input_label.endswith(':'):
                        show_row = False

            self.controls_table.setRowHidden(row, not show_row)

        # Update status bar
        visible_rows = sum(1 for row in range(self.controls_table.rowCount())
                          if not self.controls_table.isRowHidden(row))
        total_rows = self.controls_table.rowCount()
        self.statusBar().showMessage(f"Showing {visible_rows} of {total_rows} bindings")

    def toggle_detailed_view(self):
        """Toggle between default and detailed view"""
        if not self.current_profile:
            return

        # Check if user is currently editing a cell
        current_item = self.controls_table.currentItem()
        if current_item and current_item.column() == 2:  # Action (Override) column
            # Check if there's an active editor widget
            editor = self.controls_table.cellWidget(current_item.row(), current_item.column())
            if not editor:
                # Check if the item is in edit mode via QTableWidget's state
                state = self.controls_table.state()
                from PyQt6.QtWidgets import QAbstractItemView
                if state == QAbstractItemView.State.EditingState:
                    # There's an active edit - commit it by setting focus elsewhere
                    # This will trigger the itemChanged signal with the current editor value
                    self.controls_table.setCurrentItem(None)
                    # Small delay to let the edit complete
                    from PyQt6.QtCore import QTimer
                    QTimer.singleShot(0, self._complete_toggle_view)
                    return

        # No active edit, proceed immediately
        self._complete_toggle_view()

    def _complete_toggle_view(self):
        """Complete the view toggle after any pending edits are committed"""
        # Repopulate the table with the new view mode
        self.populate_controls_table()

        # Reapply filters
        self.apply_filters()

        # Update status message
        view_mode = "detailed" if self.show_detailed_checkbox.isChecked() else "default"
        self.statusBar().showMessage(f"Switched to {view_mode} view")

    def clear_filters(self):
        """Clear all filters"""
        self.search_box.clear()
        self.device_filter.setCurrentIndex(0)
        self.actionmap_filter.setCurrentIndex(0)
        self.hide_unmapped_checkbox.setChecked(False)
        self.statusBar().showMessage(f"Filters cleared - showing all {self.controls_table.rowCount()} bindings")

    def on_tab_changed(self, index: int):
        """Handle tab change - sync device selection when switching to Device View tab"""
        # Get the currently filtered device from the control table
        filtered_device = self.device_filter.currentText()

        if index == 1 and self.pdf_device_widget:
            # Device View tab
            if filtered_device and filtered_device != "All Devices":
                self.pdf_device_widget.select_device_by_name(filtered_device)

    def show_table_context_menu(self, position):
        """Show context menu for table row"""
        # Get the item at the clicked position
        item = self.controls_table.itemAt(position)
        if not item:
            return

        # Get binding data from the item
        binding_data = item.data(Qt.ItemDataRole.UserRole)
        if not binding_data:
            return

        action_map_name, binding = binding_data

        # Create context menu
        from PyQt6.QtWidgets import QMenu
        menu = QMenu(self)

        # Add "Remap Button..." action
        remap_action = menu.addAction("🎮 Remap Button...")
        remap_action.triggered.connect(lambda: self.show_remap_dialog(binding))

        # Add "Edit Label..." action
        edit_label_action = menu.addAction("✏️ Edit Label...")
        edit_label_action.triggered.connect(lambda: self.edit_label_from_context(binding))

        # Show menu at cursor position
        menu.exec(self.controls_table.viewport().mapToGlobal(position))

    def show_remap_dialog(self, binding):
        """Show the remapping dialog for a binding"""
        if not self.current_profile:
            return

        from gui.remap_dialog import RemapDialog

        dialog = RemapDialog(binding, self.current_profile, self)
        dialog.binding_changed.connect(
            lambda new_input, new_label, new_mode: self.on_binding_remapped(
                binding, new_input, new_label, new_mode
            )
        )
        dialog.exec()

    def on_binding_remapped(self, binding, new_input_code: str, new_label: str, new_activation_mode: str):
        """Handle binding changes from the remapping dialog"""
        # Update binding's input code
        binding.input_code = new_input_code

        # Update activation mode
        binding.activation_mode = new_activation_mode if new_activation_mode else None

        # Update custom label
        if new_label:
            try:
                from utils.label_overrides import get_override_manager
                override_manager = get_override_manager()

                # Get default label
                default_label = LabelGenerator.generate_action_label(binding.action_name)

                if new_label == default_label:
                    # Remove custom override
                    override_manager.remove_custom_override(binding.action_name)
                    binding.custom_label = None
                else:
                    # Save custom override
                    override_manager.set_custom_override(binding.action_name, new_label)
                    binding.custom_label = new_label

                override_manager.reload()

            except Exception as e:
                logger.error(f"Error updating label: {e}", exc_info=True)

        # Mark profile as modified
        if self.current_profile:
            self.current_profile.mark_modified()

        # Refresh UI
        self.on_profile_modified()

        logger.info(f"Remapped binding: {binding.action_name} to {new_input_code}")

    def edit_label_from_context(self, binding):
        """Edit label from context menu (simpler than full remap dialog)"""
        from PyQt6.QtWidgets import QInputDialog

        current_label = LabelGenerator.get_action_label(binding.action_name, binding)

        new_label, ok = QInputDialog.getText(
            self,
            "Edit Control Label",
            f"Edit label for {binding.action_name}:",
            text=current_label
        )

        if ok and new_label != current_label:
            try:
                from utils.label_overrides import get_override_manager
                override_manager = get_override_manager()

                # Get default label
                default_label = LabelGenerator.generate_action_label(binding.action_name)

                if new_label == default_label or not new_label:
                    # Remove custom override
                    override_manager.remove_custom_override(binding.action_name)
                    binding.custom_label = None
                else:
                    # Save custom override
                    override_manager.set_custom_override(binding.action_name, new_label)
                    binding.custom_label = new_label

                override_manager.reload()

                # Mark profile as modified (for label changes)
                if self.current_profile:
                    self.current_profile.mark_modified()

                # Refresh UI
                self.on_profile_modified()

                logger.info(f"Updated label for {binding.action_name}: {new_label}")

            except Exception as e:
                logger.error(f"Error updating label: {e}", exc_info=True)
                QMessageBox.warning(self, "Error", f"Failed to update label: {str(e)}")

    def on_item_double_clicked(self, item):
        """Handle double-click - prepare for editing"""
        if item.column() == 2:  # Action (Override) column
            # Get the binding data to store the default label for comparison
            binding_data = item.data(Qt.ItemDataRole.UserRole)
            if binding_data:
                action_map_name, binding = binding_data
                # Store the DEFAULT label (without ANY custom override - not from file, not from binding)
                # We need to get the true default: global override or auto-generated
                # To do this, we use use_override=False to skip the override manager
                self._editing_default_text = LabelGenerator.generate_action_label(binding.action_name)

                # Check if there's a global override (not custom)
                try:
                    from utils.label_overrides import get_override_manager
                except ImportError:
                    from ..utils.label_overrides import get_override_manager

                override_manager = get_override_manager()
                global_overrides = override_manager.load_global_overrides()
                if binding.action_name in global_overrides:
                    self._editing_default_text = global_overrides[binding.action_name]

                logger.debug(f"on_item_double_clicked: default='{self._editing_default_text}', current='{item.text()}'")
            else:
                self._editing_default_text = None

    def on_cell_edited(self, item):
        """Handle cell editing (only column 2: Action Override is editable)"""
        if item.column() != 2:  # Only Action (Override) column is editable
            return

        # Get the new text from the editor
        new_label = item.text().strip()

        # Get stored default text
        default_text = self._editing_default_text

        # Clear editing state
        self._editing_item = None
        self._editing_default_text = None

        # Get the binding data stored in the item
        binding_data = item.data(Qt.ItemDataRole.UserRole)
        if not binding_data:
            self.statusBar().showMessage("Error: Could not save label override")
            return

        action_map_name, binding = binding_data

        # Debug logging
        logger.debug(f"on_cell_edited: action={binding.action_name}, new_label='{new_label}', default='{default_text}'")

        # Check if user deleted everything OR if the new text equals the default
        # In both cases, remove the custom override
        if not new_label or new_label == default_text:
            # Clear the custom_label field to fall back to default
            binding.custom_label = None

            # Get the fallback label
            fallback_label = LabelGenerator.get_action_label(binding.action_name, binding)
            logger.debug(f"Label cleared or matches default, falling back to: '{fallback_label}'")

            # Update the table to show the fallback label
            self.controls_table.blockSignals(True)
            item.setText(fallback_label)
            self.controls_table.blockSignals(False)

            # Remove from the override file
            try:
                # Import with error handling for different execution contexts
                try:
                    from utils.label_overrides import get_override_manager
                except ImportError:
                    from ..utils.label_overrides import get_override_manager

                override_manager = get_override_manager()
                override_manager.remove_custom_override(binding.action_name)

                if not new_label:
                    self.statusBar().showMessage(f"Custom label removed for '{binding.action_name}' - using default: '{fallback_label}'")
                else:
                    self.statusBar().showMessage(f"Label matches default for '{binding.action_name}' - no custom override needed")
                logger.info(f"Removed custom override for '{binding.action_name}', using default: '{fallback_label}'")
            except Exception as e:
                self.statusBar().showMessage(f"Warning: Label restored but couldn't update override file: {str(e)}")
                logger.error(f"Error removing label override: {e}", exc_info=True)

            # Update Device View widget
            if self.pdf_device_widget:
                self.pdf_device_widget.load_profile(self.current_profile)

            # Force reload override manager cache and repopulate the table
            override_manager.reload()
            self.populate_controls_table()
            self.apply_filters()
            return

        # New label is different from default - save as custom override
        try:
            # Import with error handling for different execution contexts
            try:
                from utils.label_overrides import get_override_manager
            except ImportError:
                from ..utils.label_overrides import get_override_manager

            override_manager = get_override_manager()
            override_manager.set_custom_override(binding.action_name, new_label)

            # Update binding's custom_label field
            binding.custom_label = new_label

            # Update the table item to show the new custom label
            self.controls_table.blockSignals(True)
            item.setText(new_label)
            self.controls_table.blockSignals(False)

            # Update Device View widget
            if self.pdf_device_widget:
                self.pdf_device_widget.load_profile(self.current_profile)

            self.statusBar().showMessage(f"Custom label saved: '{binding.action_name}' → '{new_label}'")
            logger.info(f"Saved custom override for '{binding.action_name}': '{new_label}'")

            # Force reload override manager cache and repopulate the table
            override_manager.reload()
            self.populate_controls_table()
            self.apply_filters()
        except Exception as e:
            self.statusBar().showMessage(f"Error saving label override: {str(e)}")
            logger.error(f"Error saving label override: {e}", exc_info=True)

    def export_csv(self):
        """Export profile to CSV"""
        if not self.current_profile:
            return

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Export to CSV",
            f"{self.current_profile.profile_name}.csv",
            "CSV Files (*.csv)"
        )

        if not file_path:
            return

        try:
            import csv
            with open(file_path, 'w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)

                # Determine columns based on detailed view setting
                is_detailed = self.show_detailed_checkbox.isChecked()

                if is_detailed:
                    # Write header for detailed view (all columns)
                    writer.writerow(["Action Map", "Action (Original)", "Action (Override)", "Input Code", "Input Label", "Device"])
                    visible_cols = [0, 1, 2, 3, 4, 5]
                else:
                    # Write header for default view (4 columns)
                    writer.writerow(["Action Map", "Action", "Input Label", "Device"])
                    visible_cols = [0, 2, 4, 5]  # Action Map, Action (Override), Input Label, Device

                # Write visible rows only
                for row in range(self.controls_table.rowCount()):
                    if not self.controls_table.isRowHidden(row):
                        row_data = []
                        for col in visible_cols:
                            item = self.controls_table.item(row, col)
                            row_data.append(item.text() if item else "")
                        writer.writerow(row_data)

            QMessageBox.information(self, "Success", f"Profile exported to:\n{file_path}")
            self.statusBar().showMessage(f"Exported to CSV: {file_path}")

        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export CSV:\n{str(e)}")
            self.statusBar().showMessage("CSV export failed")

    def export_pdf(self):
        """Export profile to PDF"""
        if not self.current_profile:
            return

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Export to PDF",
            f"{self.current_profile.profile_name}.pdf",
            "PDF Files (*.pdf)"
        )

        if not file_path:
            return

        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter, landscape
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import inch

            doc = SimpleDocTemplate(file_path, pagesize=landscape(letter))
            elements = []
            styles = getSampleStyleSheet()

            # Title
            title = Paragraph(f"<b>{self.current_profile.profile_name}</b>", styles['Title'])
            elements.append(title)
            elements.append(Spacer(1, 0.2 * inch))

            # Profile info
            info_text = f"Devices: {len(self.current_profile.devices)} | "
            info_text += f"Bindings: {len(self.current_profile.get_all_bindings())}"
            info = Paragraph(info_text, styles['Normal'])
            elements.append(info)
            elements.append(Spacer(1, 0.3 * inch))

            # Determine columns based on detailed view setting
            is_detailed = self.show_detailed_checkbox.isChecked()

            if is_detailed:
                # Detailed view header and columns
                table_data = [["Action Map", "Action (Original)", "Action (Override)", "Input Code", "Input Label", "Device"]]
                visible_cols = [0, 1, 2, 3, 4, 5]
                col_widths = [1.3*inch, 1.5*inch, 1.5*inch, 1*inch, 1.5*inch, 1.2*inch]
            else:
                # Default view header and columns
                table_data = [["Action Map", "Action", "Input Label", "Device"]]
                visible_cols = [0, 2, 4, 5]
                col_widths = [2*inch, 3*inch, 2.5*inch, 1.5*inch]

            for row in range(self.controls_table.rowCount()):
                if not self.controls_table.isRowHidden(row):
                    row_data = []
                    for col in visible_cols:
                        item = self.controls_table.item(row, col)
                        row_data.append(item.text() if item else "")
                    table_data.append(row_data)

            # Create table
            table = Table(table_data, colWidths=col_widths)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
            ]))

            elements.append(table)

            # Build PDF
            doc.build(elements)

            QMessageBox.information(self, "Success", f"Profile exported to:\n{file_path}")
            self.statusBar().showMessage(f"Exported to PDF: {file_path}")

        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export PDF:\n{str(e)}")
            self.statusBar().showMessage("PDF export failed")

    def export_word(self):
        """Export profile to Word document"""
        if not self.current_profile:
            return

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Export to Word",
            f"{self.current_profile.profile_name}.docx",
            "Word Documents (*.docx)"
        )

        if not file_path:
            return

        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Title
            title = doc.add_heading(self.current_profile.profile_name, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Profile info
            info = doc.add_paragraph()
            info.add_run(f"Devices: {len(self.current_profile.devices)} | ")
            info.add_run(f"Total Bindings: {len(self.current_profile.get_all_bindings())}")
            info.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()  # Spacer

            # Device list
            doc.add_heading("Devices", 2)
            for device in self.current_profile.devices:
                device_name = device.product_name if device.product_name else device.device_type.capitalize()
                doc.add_paragraph(
                    f"{device.device_type.capitalize()} {device.instance}: {device_name}",
                    style='List Bullet'
                )

            doc.add_paragraph()  # Spacer

            # Control bindings table
            doc.add_heading("Control Bindings", 2)

            # Determine columns based on detailed view setting
            is_detailed = self.show_detailed_checkbox.isChecked()

            if is_detailed:
                headers = ["Action Map", "Action (Original)", "Action (Override)", "Input Code", "Input Label", "Device"]
                visible_cols = [0, 1, 2, 3, 4, 5]
            else:
                headers = ["Action Map", "Action", "Input Label", "Device"]
                visible_cols = [0, 2, 4, 5]

            # Create table
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Light Grid Accent 1'

            # Header row
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                # Make header bold
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Add data rows
            for row in range(self.controls_table.rowCount()):
                if not self.controls_table.isRowHidden(row):
                    row_cells = table.add_row().cells
                    for i, col in enumerate(visible_cols):
                        item = self.controls_table.item(row, col)
                        row_cells[i].text = item.text() if item else ""

            # Save document
            doc.save(file_path)

            QMessageBox.information(self, "Success", f"Profile exported to:\n{file_path}")
            self.statusBar().showMessage(f"Exported to Word: {file_path}")

        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export Word document:\n{str(e)}")
            self.statusBar().showMessage("Word export failed")

    def export_graphic(self):
        """Export device graphic - delegate to Device View widget"""
        # Check if Device View tab is active
        current_tab = self.tab_widget.currentIndex()
        if current_tab == 1 and self.pdf_device_widget:
            # Device View tab
            self.pdf_device_widget.export_graphic()
        else:
            # Not on Device View tab
            QMessageBox.warning(self, "Export Graphic", "Please switch to the Device View tab first.")

    def on_profile_modified(self):
        """Handle profile modification - update UI to reflect changes"""
        if not self.current_profile:
            return

        # Update window title with modified indicator
        base_title = f"Star Citizen Profile Editor v{self.version}"
        if self.current_profile.is_modified:
            self.setWindowTitle(f"{base_title} - {self.current_profile.profile_name} *")
            # Enable save button when modified
            self.save_profile_btn.setEnabled(True)
        else:
            self.setWindowTitle(f"{base_title} - {self.current_profile.profile_name}")
            # Disable save button when not modified
            self.save_profile_btn.setEnabled(False)

        # Refresh the controls table
        self.populate_controls_table()
        self.apply_filters()

        # Refresh device view if active
        if self.pdf_device_widget and self.pdf_device_widget.current_device:
            self.pdf_device_widget.load_device_pdf()

    def save_profile(self):
        """Save the modified profile to XML"""
        if not self.current_profile or not self.current_profile.is_modified:
            return

        # Get output path - use default naming with _scpe suffix
        output_path = None

        if self.current_profile.source_xml_path:
            # Generate default filename with _scpe suffix
            source_path = self.current_profile.source_xml_path
            base_name = os.path.splitext(os.path.basename(source_path))[0]

            # Remove _exported suffix if present
            if base_name.endswith('_exported'):
                base_name = base_name[:-9]

            suggested_name = f"{base_name}_scpe.xml"
            suggested_dir = os.path.dirname(source_path)
            default_path = os.path.join(suggested_dir, suggested_name)
        else:
            # No source path, use profile name
            suggested_name = f"{self.current_profile.profile_name}_scpe.xml"
            default_path = suggested_name

        # Show save dialog with default name
        output_path, _ = QFileDialog.getSaveFileName(
            self,
            "Export Profile",
            default_path,
            "XML Files (*.xml);;All Files (*)"
        )

        if not output_path:
            return

        # Export profile to XML
        try:
            from exporters.xml_exporter import export_profile

            # Use original source path if available, otherwise create from scratch
            source_path = self.current_profile.source_xml_path

            if source_path and os.path.exists(source_path):
                success = export_profile(self.current_profile, source_path, output_path)
            else:
                # Create new profile from scratch
                from exporters.xml_exporter import XMLExporter
                success = XMLExporter.create_new_profile(self.current_profile, output_path)

            if success:
                # Mark as saved
                self.current_profile.mark_saved()
                self.current_profile.source_xml_path = output_path

                # Update UI
                self.on_profile_modified()

                # Show success message
                QMessageBox.information(
                    self,
                    "Success",
                    f"Profile saved successfully!\n\nSaved to:\n{output_path}"
                )

                self.statusBar().showMessage(f"Profile saved: {os.path.basename(output_path)}")
                logger.info(f"Saved profile to: {output_path}")

            else:
                QMessageBox.critical(
                    self,
                    "Save Error",
                    "Failed to save profile. Check the log for details."
                )

        except Exception as e:
            logger.error(f"Error saving profile: {e}", exc_info=True)
            QMessageBox.critical(
                self,
                "Save Error",
                f"Failed to save profile:\n{str(e)}"
            )

    def show_help(self):
        """Show the user guide in a dialog"""
        # Create help dialog
        help_dialog = QDialog(self)
        help_dialog.setWindowTitle(f"User Guide - Star Citizen Profile Editor v{self.version}")
        help_dialog.setGeometry(100, 100, 900, 700)

        # Layout
        layout = QVBoxLayout()
        help_dialog.setLayout(layout)

        # Text browser for displaying markdown as HTML
        browser = QTextBrowser()
        browser.setOpenExternalLinks(True)

        # Load the README.md file
        user_guide_path = os.path.join(
            os.path.dirname(os.path.dirname(os.path.dirname(__file__))),
            "README.md"
        )

        try:
            with open(user_guide_path, 'r', encoding='utf-8') as f:
                markdown_content = f.read()

            # Convert markdown to HTML (basic conversion)
            html_content = self.markdown_to_html(markdown_content)
            browser.setHtml(html_content)

        except FileNotFoundError:
            browser.setHtml("<h1>User Guide Not Found</h1><p>The README.md file could not be found.</p>")
        except Exception as e:
            browser.setHtml(f"<h1>Error Loading Guide</h1><p>{str(e)}</p>")

        layout.addWidget(browser)

        # Close button
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(help_dialog.close)
        layout.addWidget(close_btn)

        help_dialog.exec()

    def markdown_to_html(self, markdown_text):
        """Convert markdown to HTML (basic implementation)"""
        html = "<html><head><style>"
        html += "body { font-family: Arial, sans-serif; line-height: 1.8; padding: 20px; font-size: 14px; }"
        html += "h1 { color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 10px; font-size: 32px; font-weight: bold; margin-top: 20px; }"
        html += "h2 { color: #34495e; border-bottom: 2px solid #95a5a6; padding-bottom: 5px; margin-top: 30px; font-size: 26px; font-weight: bold; }"
        html += "h3 { color: #555; margin-top: 20px; font-size: 20px; font-weight: bold; }"
        html += "h4 { color: #666; margin-top: 15px; font-size: 16px; font-weight: bold; }"
        html += "p { font-size: 14px; margin: 10px 0; }"
        html += "li { font-size: 14px; margin: 5px 0; }"
        html += "a { color: #3498db; text-decoration: none; }"
        html += "a:hover { text-decoration: underline; }"
        html += "code { background-color: #f4f4f4; padding: 2px 6px; border-radius: 3px; font-family: 'Courier New', monospace; font-size: 13px; }"
        html += "pre { background-color: #f4f4f4; padding: 10px; border-radius: 5px; overflow-x: auto; font-size: 13px; }"
        html += "ul { margin-left: 20px; font-size: 14px; }"
        html += "ol { margin-left: 20px; font-size: 14px; }"
        html += "strong { color: #2c3e50; font-weight: bold; }"
        html += "blockquote { border-left: 4px solid #3498db; padding-left: 15px; color: #555; font-style: italic; font-size: 14px; }"
        html += "table { border-collapse: collapse; width: 100%; margin: 20px 0; }"
        html += "th, td { border: 1px solid #ddd; padding: 8px; text-align: left; font-size: 14px; }"
        html += "th { background-color: #3498db; color: white; font-weight: bold; }"
        html += "tr:nth-child(even) { background-color: #f2f2f2; }"
        html += "</style></head><body>"

        lines = markdown_text.split('\n')
        in_code_block = False
        in_list = False

        for line in lines:
            # Code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    html += "</pre>"
                    in_code_block = False
                else:
                    html += "<pre><code>"
                    in_code_block = True
                continue

            if in_code_block:
                html += line + "\n"
                continue

            # Headers (with anchor IDs for links)
            if line.startswith('# '):
                header_text = line[2:].strip()
                anchor_id = self.create_anchor_id(header_text)
                html += f"<h1 id='{anchor_id}'>{header_text}</h1>"
            elif line.startswith('## '):
                header_text = line[3:].strip()
                anchor_id = self.create_anchor_id(header_text)
                html += f"<h2 id='{anchor_id}'>{header_text}</h2>"
            elif line.startswith('### '):
                header_text = line[4:].strip()
                anchor_id = self.create_anchor_id(header_text)
                html += f"<h3 id='{anchor_id}'>{header_text}</h3>"
            elif line.startswith('#### '):
                header_text = line[5:].strip()
                anchor_id = self.create_anchor_id(header_text)
                html += f"<h4 id='{anchor_id}'>{header_text}</h4>"
            # Ordered lists
            elif line.strip().startswith(tuple(f"{i}. " for i in range(1, 100))):
                if not in_list:
                    html += "<ol>"
                    in_list = 'ol'
                # Extract the list item text (after the number and period)
                item_text = line.strip().split('. ', 1)[1] if '. ' in line else line.strip()
                html += f"<li>{self.format_inline_markdown(item_text)}</li>"
            # Unordered lists
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                if in_list != 'ul':
                    if in_list:
                        html += f"</{in_list}>"
                    html += "<ul>"
                    in_list = 'ul'
                html += f"<li>{self.format_inline_markdown(line.strip()[2:])}</li>"
            # Horizontal rule
            elif line.strip() == '---':
                if in_list:
                    html += f"</{in_list}>"
                    in_list = False
                html += "<hr/>"
            # Empty line
            elif not line.strip():
                if in_list:
                    html += f"</{in_list}>"
                    in_list = False
                html += "<br/>"
            # Regular paragraph
            else:
                if in_list:
                    html += f"</{in_list}>"
                    in_list = False
                html += f"<p>{self.format_inline_markdown(line)}</p>"

        if in_list:
            html += f"</{in_list}>"

        html += "</body></html>"
        return html

    def create_anchor_id(self, header_text):
        """Create an anchor ID from header text (GitHub-style)"""
        import re
        # Convert to lowercase
        anchor = header_text.lower()
        # Replace spaces with hyphens
        anchor = anchor.replace(' ', '-')
        # Remove special characters except hyphens
        anchor = re.sub(r'[^\w\-]', '', anchor)
        return anchor

    def format_inline_markdown(self, text):
        """Format inline markdown (bold, italic, code, links)"""
        import re

        # Bold
        text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
        # Italic
        text = re.sub(r'\*(.+?)\*', r'<em>\1</em>', text)
        # Inline code
        text = re.sub(r'`(.+?)`', r'<code>\1</code>', text)
        # Links (handle both URLs and anchors)
        text = re.sub(r'\[(.+?)\]\((.+?)\)', r'<a href="\2">\1</a>', text)

        return text

    def open_discord_link(self, event):
        """Open Discord invite link in browser"""
        discord_url = "https://discord.gg/BNzRegKZ7k"
        QDesktopServices.openUrl(QUrl(discord_url))

    def open_paypal_donation(self, event):
        """Open PayPal donation link in browser"""
        paypal_url = "https://paypal.me/RighteousKill"
        QDesktopServices.openUrl(QUrl(paypal_url))

    def open_venmo_donation(self, event):
        """Open Venmo donation link in browser"""
        venmo_url = "https://venmo.com/u/Amr-Abouelleil"
        QDesktopServices.openUrl(QUrl(venmo_url))


if __name__ == "__main__":
    from PyQt6.QtWidgets import QApplication

    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())